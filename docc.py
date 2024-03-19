from docx import Document
from docx.shared import Inches
from PIL import Image

def add_image_to_header(header, image_path):
    header_paragraph = header.paragraphs[0]
    run = header_paragraph.add_run()
    run.add_picture(image_path, width=Inches(6.5))  # Adjust width as needed

def add_image_to_footer(footer, image_path):
    footer_paragraph = footer.paragraphs[0]
    run = footer_paragraph.add_run()
    run.add_picture(image_path, width=Inches(6.5))  # Adjust width as needed

def create_word_document(content, header_image_path, footer_image_path):
    doc = Document()

    # Add header with image
    header = doc.sections[0].header
    add_image_to_header(header, header_image_path)

    # Add footer with image
    footer = doc.sections[0].footer
    add_image_to_footer(footer, footer_image_path)

    # Add your content with newline replacements
    for line in content.split('\n'):
        doc.add_paragraph(line)

    # Save the document
    doc.save('your_document.docx')

# Example usage:
#your_string_with_newlines = "03/17/2024\n{name of Judge}\nFULTON CITY COURT\nMUNI BLDG 141 S.1ST ST.\nFULTON, NY, 13069\n\nRe:\tMAWSON ASHLEY  (d.o.b.  08/08/1991)   \nSPEED IN ZONE  (NYS V AND T LAW § 1180D)\n\t\tUniform Traffic Ticket # 0999HRQKJB\n{Description of violation}  ({in violation of} § {Violation Section})\n\t\tUniform Traffic Ticket # 0999HRQKJB\n{Description of violation}  ({in violation of} § {Violation Section})\n\t\tUniform Traffic Ticket # 0999HRQKJB\n\tViolation Date: 02/28/2024\nDear Judge:\nEnclosed is an offer of disposition from the prosecutor recommending that the defendant plead guilty to a SPEED IN ZONE  In violation of NYS V AND T LAW § 1180D  in satisfaction of the above charge(s). \nPlease notify me if said disposition is satisfactory and, if so, the amount of the fine, and I will see that this matter is disposed of as soon as possible.  Thank you for your consideration. \nVery truly yours,\n\nThomas B. Mafrici"
  # Replace with the path to your footer image

